"""
Script tạo báo cáo kỹ thuật Word cho hệ thống điều khiển VFD bằng ESP32-S3.
Yêu cầu: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Đặt màu nền cho ô bảng."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    heading.paragraph_format.space_before = Pt(14) if level == 1 else Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    if indent:
        p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    return p


def add_formula(doc, formula_text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(formula_text)
    run.font.size = Pt(12)
    run.font.name = 'Courier New'
    run.bold = True
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"▶  {text}")
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x17, 0x5F, 0x1E)
    return p


def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.75)
    if bold_part and text.startswith(bold_part):
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(12)
        rest = p.add_run(text[len(bold_part):])
        rest.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.size = Pt(12)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], '2E74B5')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = 'F2F7FC' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            set_cell_bg(row_cells[ci], bg)
            run = row_cells[ci].paragraphs[0].runs
            if run:
                run[0].font.size = Pt(11)

    if col_widths:
        for i, col in enumerate(table.columns):
            col.width = Cm(col_widths[i])
    return table


def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('─' * 70)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ─────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────

doc = Document()

# Lề
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

# Font mặc định
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ─────────────────────────────────────────────
# TRANG BÌA
# ─────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run('BÁO CÁO KỸ THUẬT')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle_para.add_run(
    'Hệ thống điều khiển tốc độ động cơ qua Biến Tần (VFD)\n'
    'sử dụng ESP32-S3, Cảm biến siêu âm và Modbus RTU'
)
r2.bold = True
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = info_para.add_run(
    f'Ngày lập: {datetime.date.today().strftime("%d/%m/%Y")}\n'
    'Nền tảng: ESP32-S3 | PlatformIO | Arduino Framework\n'
    'Giao tiếp: RS485 + Modbus RTU\n'
    'Màn hình: OLED SH1106\n'
    'Bộ nhớ: NVS Flash'
)
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_page_break()

# ─────────────────────────────────────────────
# MỤC LỤC (tĩnh)
# ─────────────────────────────────────────────

add_heading(doc, 'MỤC LỤC', level=1)
toc_items = [
    ('1.', 'ABSTRACT (Tóm tắt)'),
    ('2.', 'INTRODUCTION (Giới thiệu)'),
    ('3.', 'SYSTEM OVERVIEW (Tổng quan hệ thống)'),
    ('4.', 'HARDWARE DESIGN (Thiết kế phần cứng)'),
    ('5.', 'SOFTWARE ARCHITECTURE (Kiến trúc phần mềm)'),
    ('6.', 'CONTROL ALGORITHM (Thuật toán điều khiển)'),
    ('  6.1', 'Đọc khoảng cách'),
    ('  6.2', 'Mapping khoảng cách → tần số'),
    ('  6.3', 'Logic BOOST'),
    ('  6.4', 'Anti-spam Modbus'),
    ('  6.5', 'Multi-task FreeRTOS'),
    ('7.', 'COMMUNICATION PROTOCOL (Giao thức truyền thông)'),
    ('8.', 'DATA STORAGE (Lưu cấu hình)'),
    ('9.', 'USER INTERFACE (Giao diện người dùng)'),
    ('10.', 'SYSTEM MODES (Chế độ vận hành)'),
    ('11.', 'PERFORMANCE ANALYSIS (Phân tích hiệu suất)'),
    ('12.', 'LIMITATIONS (Giới hạn)'),
    ('13.', 'FUTURE IMPROVEMENTS (Cải tiến tương lai)'),
    ('14.', 'CONCLUSION (Kết luận)'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5) if num.startswith('  ') else Cm(0)
    r = p.add_run(f'{num.strip()}    {title}')
    r.font.size = Pt(12)

doc.add_page_break()

# ─────────────────────────────────────────────
# CHƯƠNG 1 – ABSTRACT
# ─────────────────────────────────────────────

add_heading(doc, '1. ABSTRACT (Tóm tắt)')

add_paragraph(doc,
    'Đề tài này xây dựng một hệ thống điều khiển tốc độ động cơ thông qua biến tần '
    'nhằm tự động hóa quá trình cấp dây theo trạng thái thực tế của cuộn dây trong '
    'sản xuất. Hệ thống sử dụng vi điều khiển ESP32-S3 làm bộ xử lý trung tâm, cảm '
    'biến siêu âm để đo khoảng cách đến drum, giao tiếp RS485 kết hợp Modbus RTU để '
    'gửi lệnh điều khiển đến biến tần, đồng thời hiển thị trạng thái vận hành trên '
    'màn hình OLED SH1106.'
)
add_paragraph(doc,
    'Mục tiêu chính của hệ thống là chuyển từ mô hình điều khiển thủ công phụ thuộc '
    'kinh nghiệm người vận hành sang mô hình điều khiển tự động có quy luật, ổn định '
    'và dễ lặp lại. Khi lượng dây thay đổi, khoảng cách đo được cũng thay đổi tương '
    'ứng; từ đó bộ điều khiển tự động tính toán tần số phù hợp để điều chỉnh tốc độ motor.'
)

add_paragraph(doc, 'Ý nghĩa cải tiến thể hiện ở ba điểm nổi bật:', bold=True)
add_bullet(doc, 'Giảm phụ thuộc vào thao tác cảm tính của operator.')
add_bullet(doc, 'Tăng tính ổn định của quá trình cấp dây, hạn chế thiếu hoặc dư dây cục bộ.')
add_bullet(doc, 'Tạo nền tảng để mở rộng lên các mô hình điều khiển thông minh hơn như PID, giám sát từ xa.')

add_note(doc,
    'Nói dễ hiểu: trước đây công nhân phải "nhìn cuộn dây rồi tự vặn tốc độ", '
    'còn bây giờ ESP32 làm việc đó tự động, đều và nhanh hơn.'
)

# ─────────────────────────────────────────────
# CHƯƠNG 2 – INTRODUCTION
# ─────────────────────────────────────────────

add_heading(doc, '2. INTRODUCTION (Giới thiệu)')

add_paragraph(doc,
    'Trong nhiều dây chuyền công nghiệp, đặc biệt là các hệ thống quấn dây, nhả dây '
    'hoặc cấp vật liệu dạng cuộn, việc duy trì tốc độ motor phù hợp với trạng thái '
    'drum là yêu cầu rất quan trọng. Nếu motor quay quá chậm khi drum đang thiếu dây, '
    'dây sẽ bị căng không ổn định, ảnh hưởng chất lượng sản phẩm. Ngược lại, nếu quay '
    'quá nhanh khi drum đã đầy, hệ thống có thể bị chùng dây, rối dây hoặc rung giật cơ khí.'
)
add_paragraph(doc,
    'Thực tế tại nhiều nhà máy, việc điều chỉnh tốc độ biến tần vẫn dựa vào thao tác '
    'thủ công. Người vận hành quan sát trạng thái drum rồi tăng, giảm tốc độ theo kinh nghiệm. '
    'Phương pháp này tồn tại nhiều nhược điểm: phụ thuộc tay nghề, phản ứng chậm, khó đồng '
    'nhất giữa các ca sản xuất và khó duy trì chất lượng ổn định trong thời gian dài.'
)
add_paragraph(doc,
    'Xuất phát từ bài toán đó, đề tài đặt mục tiêu thiết kế một bộ điều khiển nhúng '
    'nhỏ gọn, chi phí hợp lý nhưng có khả năng làm việc thực tế trong môi trường công nghiệp. '
    'Hệ thống phải đo được trạng thái drum, chuyển đổi thành tốc độ phù hợp, giao tiếp được '
    'với biến tần qua Modbus RTU, có cơ chế tăng tốc bù khi thiếu dây, lưu tham số vào Flash '
    'và hỗ trợ chế độ AUTO, MANUAL, TEST.'
)

# ─────────────────────────────────────────────
# CHƯƠNG 3 – SYSTEM OVERVIEW
# ─────────────────────────────────────────────

add_heading(doc, '3. SYSTEM OVERVIEW (Tổng quan hệ thống)')

add_heading(doc, '3.1 Mô tả tổng thể', level=2)
add_paragraph(doc,
    'Hệ thống được tổ chức theo kiến trúc cảm biến – xử lý – truyền thông – chấp hành. '
    'Luồng hoạt động tổng quát được mô tả như sau:'
)
add_formula(doc,
    'Cảm biến siêu âm  →  ESP32-S3  →  Thuật toán điều khiển  →  Modbus RTU  →  RS485  →  VFD  →  Motor'
)

add_heading(doc, '3.2 Sơ đồ khối hệ thống', level=2)
add_paragraph(doc, 'Hệ thống bao gồm các khối chức năng chính sau:')
add_bullet(doc, 'Khối đo lường: cảm biến siêu âm đo khoảng cách, cảm biến số phát hiện thiếu dây.')
add_bullet(doc, 'Khối điều khiển trung tâm: ESP32-S3 đọc tín hiệu, lọc nhiễu, tính tần số đích.')
add_bullet(doc, 'Khối truyền thông công nghiệp: UART → Module RS485 → Modbus RTU → VFD.')
add_bullet(doc, 'Khối chấp hành: VFD nhận lệnh tần số và run/stop, motor thay đổi tốc độ.')
add_bullet(doc, 'Khối giao diện: OLED hiển thị trạng thái, Web HMI hỗ trợ cấu hình từ xa.')

add_heading(doc, '3.3 Vai trò từng thành phần', level=2)
add_table(doc,
    headers=['Thành phần', 'Vai trò'],
    rows=[
        ['ESP32-S3', 'Bộ não trung tâm: thu thập dữ liệu, tính toán thuật toán, gửi lệnh Modbus, quản lý OLED và Flash.'],
        ['Cảm biến siêu âm', 'Đo khoảng cách đến drum, cung cấp đại lượng đầu vào chính cho thuật toán.'],
        ['VFD (Biến tần)', 'Biến lệnh tần số từ ESP32 thành tốc độ quay thực tế của động cơ.'],
        ['OLED SH1106', 'Cung cấp phản hồi trực tiếp tại hiện trường: distance, freq, mode, boost.'],
        ['RS485 + Modbus', 'Kênh truyền thông chống nhiễu, phù hợp môi trường công nghiệp.'],
        ['Flash NVS', 'Lưu cấu hình bền vững, không mất khi mất điện.'],
    ],
    col_widths=[4.5, 12.0]
)

# ─────────────────────────────────────────────
# CHƯƠNG 4 – HARDWARE DESIGN
# ─────────────────────────────────────────────

add_heading(doc, '4. HARDWARE DESIGN (Thiết kế phần cứng)')

add_heading(doc, '4.1 Sơ đồ kết nối chi tiết', level=2)
add_table(doc,
    headers=['Tín hiệu', 'ESP32-S3 Pin', 'Kết nối đến', 'Ghi chú'],
    rows=[
        ['TRIG cảm biến', 'GPIO 4', 'TRIG của HC-SR04', 'Xung kích phát sóng siêu âm'],
        ['ECHO cảm biến', 'GPIO 5', 'ECHO của HC-SR04', 'Nhận thời gian phản hồi'],
        ['UART TX (RS485)', 'GPIO 17', 'DI của module RS485', 'Gửi lệnh Modbus RTU'],
        ['UART RX (RS485)', 'GPIO 18', 'RO của module RS485', 'Nhận phản hồi từ VFD'],
        ['DE / RE (RS485)', 'GPIO 16', 'DE + RE module RS485', 'Điều khiển chiều truyền bán song công'],
        ['I2C SDA (OLED)', 'GPIO 8', 'SDA của SH1106', 'Dữ liệu màn hình'],
        ['I2C SCL (OLED)', 'GPIO 9', 'SCL của SH1106', 'Clock màn hình'],
        ['Cảm biến thiếu dây', 'GPIO 10 (INPUT_PULLUP)', 'NO contact của cảm biến', 'Kích hoạt BOOST logic'],
    ],
    col_widths=[4.0, 3.5, 4.5, 5.5]
)

add_heading(doc, '4.2 Giải thích nguyên lý TRIG / ECHO', level=2)
add_paragraph(doc,
    'Cảm biến siêu âm hoạt động dựa trên thời gian bay của sóng âm. '
    'Chân TRIG nhận xung kích ngắn từ ESP32, cảm biến phát chùm siêu âm. '
    'Khi sóng gặp vật cản và phản xạ về, chân ECHO lên mức cao trong khoảng '
    'thời gian tương ứng với thời gian sóng truyền đi và quay lại.'
)
add_formula(doc, 'd = (v × t) / 2       với v ≈ 343 m/s,  t tính bằng micro giây')

add_heading(doc, '4.3 Giải thích DE / RE trong RS485', level=2)
add_paragraph(doc,
    'RS485 là chuẩn bán song công — tại một thời điểm chỉ nên có một bên phát. '
    'Chân DE (Driver Enable) bật khi ESP32 muốn gửi dữ liệu. '
    'Chân RE (Receiver Enable) được kéo cùng với DE để tối giản mạch điện. '
    'Trước khi gửi gói Modbus, ESP32 bật phát (DE=HIGH); sau khi xong trả lại nhận (DE=LOW).'
)

add_heading(doc, '4.4 Lý do chọn linh kiện', level=2)
add_bullet(doc, 'ESP32-S3: hiệu năng tốt, hỗ trợ Wi-Fi, đủ ngoại vi cho cả điều khiển và HMI cục bộ.')
add_bullet(doc, 'Cảm biến siêu âm: giá thành thấp, dễ lắp, không tiếp xúc cơ khí.')
add_bullet(doc, 'RS485: chuẩn công nghiệp, chống nhiễu tốt, truyền xa, tương thích Modbus RTU.')
add_bullet(doc, 'OLED SH1106: hiển thị rõ, tiêu thụ điện thấp, đủ cho thông tin vận hành cốt lõi.')

# ─────────────────────────────────────────────
# CHƯƠNG 5 – SOFTWARE ARCHITECTURE
# ─────────────────────────────────────────────

add_heading(doc, '5. SOFTWARE ARCHITECTURE (Kiến trúc phần mềm)')

add_heading(doc, '5.1 Cấu trúc file', level=2)
add_table(doc,
    headers=['File', 'Vai trò'],
    rows=[
        ['src/Config.h', 'Cấu hình phần cứng: chân kết nối, địa chỉ I2C, baud rate, Modbus ID.'],
        ['src/Settings.h', 'Định nghĩa cấu trúc tham số hệ thống và cơ chế lưu/đọc từ Flash NVS.'],
        ['src/VFDManager.h', 'Đóng gói giao tiếp Modbus RTU với biến tần, retry logic, DE/RE control.'],
        ['src/WiFiWebManager.h/cpp', 'Wi-Fi, Web server, chia sẻ runtime data qua mutex.'],
        ['src/main.cpp', 'Điều phối toàn bộ: thuật toán điều khiển, lọc tín hiệu, OLED, FreeRTOS task.'],
    ],
    col_widths=[5.0, 12.5]
)

add_heading(doc, '5.2 Luồng chương trình', level=2)
add_paragraph(doc, 'Quy trình khởi động và vận hành:')
steps = [
    'Khởi tạo Serial, UART, OLED, Wi-Fi và mutex đồng bộ (semaphore).',
    'Đọc cấu hình đã lưu từ Flash NVS, kiểm tra hợp lệ và sanitize.',
    'Khởi tạo kết nối Modbus với VFD, gửi lệnh cho phép chạy (0x0012).',
    'Vòng lặp chính (loop) mỗi 20 ms: đọc cảm biến → lọc nhiễu → tính tần số → kiểm tra BOOST.',
    'Nếu cần gửi lệnh: tạo task FreeRTOS nền để gửi Modbus RTU, không block vòng lặp.',
    'Mỗi 100 ms: cập nhật OLED với trạng thái mới nhất.',
    'Wi-Fi serviceReconnect() được gọi mỗi chu kỳ để duy trì kết nối nền.',
]
for i, s in enumerate(steps, 1):
    add_bullet(doc, f'Bước {i}: {s}')

add_heading(doc, '5.3 Task FreeRTOS gửi Modbus', level=2)
add_paragraph(doc,
    'Một điểm kỹ thuật quan trọng là tách việc gửi Modbus sang tác vụ nền FreeRTOS. '
    'Điều này giúp vòng lặp chính không bị chặn bởi thời gian truyền nối tiếp hoặc '
    'retry trên bus RS485. Lợi ích cụ thể:'
)
add_bullet(doc, 'Vòng lặp chính giữ nhịp đọc cảm biến đều đặn.')
add_bullet(doc, 'Hiển thị OLED không bị giật hay delay.')
add_bullet(doc, 'Tránh treo logic điều khiển khi đường truyền chậm.')
add_bullet(doc, 'Cờ isNetTaskRunning đảm bảo tại một thời điểm chỉ có một task Modbus hoạt động.')

# ─────────────────────────────────────────────
# CHƯƠNG 6 – CONTROL ALGORITHM
# ─────────────────────────────────────────────

add_heading(doc, '6. CONTROL ALGORITHM (Thuật toán điều khiển)')

add_heading(doc, '6.1 Đọc khoảng cách', level=2)
add_paragraph(doc,
    'Cảm biến siêu âm đo khoảng cách bằng nguyên lý thời gian truyền sóng (TOF – Time of Flight). '
    'ESP32 phát xung kích, cảm biến phát sóng siêu âm. Khi sóng phản xạ về, chân ECHO trả về '
    'tín hiệu có độ rộng tương ứng thời gian di chuyển của sóng.'
)
add_formula(doc, 'd = (v × t) / 2       [cm]')
add_paragraph(doc,
    'Trong đó: d = khoảng cách (cm),  v = vận tốc âm ≈ 343 m/s,  t = thời gian đo (µs).'
)

add_heading(doc, '6.2 Lọc nhiễu tín hiệu', level=2)
add_paragraph(doc,
    'Tín hiệu siêu âm ngoài thực tế bị nhiễu do rung cơ khí, bề mặt không đều, góc phản xạ. '
    'Hệ thống áp dụng bộ lọc trung bình lũy tiến (Exponential Moving Average – EMA):'
)
add_formula(doc, 'd_filtered = 0.7 × d_old  +  0.3 × d_raw')
add_paragraph(doc,
    'Hệ số 0.7 ưu tiên ổn định (không nhảy đột ngột). '
    'Hệ số 0.3 đảm bảo hệ thống vẫn phản ứng được với thay đổi thực sự. '
    'Đây là lựa chọn cân bằng tốt cho hệ cơ khí có quán tính vừa phải.'
)
add_note(doc,
    'Nói đơn giản: cảm biến "nói" một con số hơi khác nhau mỗi lần; '
    'bộ điều khiển không tin hoàn toàn ngay mà lấy trung bình thông minh để quyết định.'
)

add_heading(doc, '6.3 Mapping khoảng cách → tần số', level=2)
add_paragraph(doc,
    'Đây là lõi của thuật toán điều khiển. Ý tưởng: khoảng cách càng lớn thì '
    'dây càng thiếu, motor cần quay nhanh hơn; khoảng cách càng nhỏ thì motor '
    'cần chạy chậm lại. Công thức nội suy tuyến tính:'
)
add_formula(doc,
    'f = f_min + (d - d_min) × (f_max - f_min) / (d_max - d_min)'
)
add_paragraph(doc,
    'Trong đó:\n'
    '  d_min / d_max : khoảng cách biên min/max cấu hình (cm)\n'
    '  f_min / f_max : tần số biên min/max cấu hình (Hz)'
)

add_paragraph(doc, 'Ví dụ minh họa:', bold=True)
add_paragraph(doc,
    'Cấu hình: d_min=10cm, d_max=50cm, f_min=20Hz, f_max=50Hz.',
    indent=True
)
add_formula(doc,
    'Đo được d=30cm  →  f = 20 + (30-10)×(50-20)/(50-10)  =  35 Hz'
)
add_note(doc, 'Hệ thống yêu cầu VFD chạy ở 35 Hz khi khoảng cách drum là 30 cm.')

add_heading(doc, '6.4 Logic BOOST – Tăng tốc bù khi thiếu dây', level=2)
add_paragraph(doc,
    'Khi cảm biến thiếu dây kích hoạt, hệ thống áp dụng hệ số BOOST lên tần số đã tính. '
    'Đây là cơ chế bù tốc tạm thời để bắt kịp lượng dây bị hụt.'
)
add_formula(doc,
    'f_cmd = f_map × (1 + boostfactor / 100)'
)
add_paragraph(doc, 'Ví dụ: f_map = 35 Hz, boost = 20%  →  f_cmd = 35 × 1.2 = 42 Hz', indent=True)

add_paragraph(doc, 'Hệ thống có 3 cấp BOOST leo thang theo thời gian kích hoạt:', bold=True)
add_table(doc,
    headers=['Cấp', 'Điều kiện kích hoạt', 'Phần trăm BOOST mặc định', 'Thời gian giữ mặc định'],
    rows=[
        ['Cấp 1', 'Cảm biến vừa kích', '20%', '5 giây'],
        ['Cấp 2', 'Cảm biến giữ ≥ 0.8 giây', '35%', '4 giây'],
        ['Cấp 3', 'Cảm biến giữ ≥ 1.6 giây', '50%', '3 giây'],
    ],
    col_widths=[2.0, 5.0, 5.0, 4.5]
)
add_paragraph(doc,
    'Sau khi cảm biến nhả, BOOST không tắt ngay mà giữ trong thời gian hold tương ứng cấp. '
    'Sau đó giảm dần theo cơ chế decay để tránh rung trạng thái ON/OFF làm hệ cơ khí dao động.'
)

add_heading(doc, '6.5 Anti-spam Modbus', level=2)
add_paragraph(doc,
    'Hệ thống bảo vệ bus RS485 bằng 3 lớp:'
)
add_bullet(doc, 'Giới hạn chu kỳ cập nhật: lệnh chỉ được xét gửi theo chu kỳ cố định.')
add_bullet(doc, 'Vùng chết tần số: chỉ gửi khi |f_new - f_last| ≥ 0.2 Hz.')
add_bullet(doc, 'Cờ task: tại một thời điểm chỉ có một FreeRTOS task Modbus hoạt động.')
add_note(doc, 'Kết quả: bus sạch hơn, VFD ổn định hơn, độ tin cậy tổng thể cao hơn.')

add_heading(doc, '6.6 Multi-task FreeRTOS', level=2)
add_paragraph(doc,
    'ESP32-S3 hỗ trợ FreeRTOS, phù hợp cho hệ thống vừa cần phản ứng nhanh vừa '
    'xử lý nhiều nhóm công việc khác nhau. Nếu tất cả dồn vào một vòng lặp đồng bộ, '
    'chỉ cần Modbus bị chậm là toàn hệ thống khựng. Phân tách task giải quyết vấn đề này:'
)
add_table(doc,
    headers=['Task / Nhóm công việc', 'Tần suất', 'Ghi chú'],
    rows=[
        ['Đọc cảm biến + tính toán', '~50 Hz (20 ms)', 'Trong vòng lặp chính'],
        ['Gửi lệnh Modbus', 'Khi có dữ liệu mới', 'Task nền FreeRTOS trên core 1'],
        ['Cập nhật OLED', '~10 Hz (100 ms)', 'Trong vòng lặp chính'],
        ['Wi-Fi reconnect', 'Mỗi chu kỳ loop', 'Non-blocking backoff'],
        ['Web server', 'On-demand', 'ESPAsyncWebServer'],
    ],
    col_widths=[6.0, 4.0, 7.5]
)

# ─────────────────────────────────────────────
# CHƯƠNG 7 – COMMUNICATION PROTOCOL
# ─────────────────────────────────────────────

add_heading(doc, '7. COMMUNICATION PROTOCOL (Giao thức truyền thông)')

add_heading(doc, '7.1 Modbus RTU', level=2)
add_paragraph(doc,
    'Modbus RTU là giao thức truyền thông rất phổ biến trong công nghiệp nhờ đơn giản, '
    'ổn định và dễ tích hợp với biến tần, PLC, HMI. Trong đề tài này, ESP32 đóng vai '
    'trò master, còn biến tần là slave với địa chỉ 0x01.'
)

add_heading(doc, '7.2 Thanh ghi điều khiển chính', level=2)
add_table(doc,
    headers=['Thanh ghi', 'Địa chỉ', 'Chiều', 'Chức năng'],
    rows=[
        ['Tần số đặt', '0x2001', 'Write', 'Ghi giá trị tần số × 100 (ví dụ: 35 Hz → 3500)'],
        ['Run / Stop', '0x2000', 'Write', '0x0012 = RUN, 0x0001 = STOP'],
    ],
    col_widths=[3.0, 3.0, 2.0, 9.5]
)

add_heading(doc, '7.3 Cách đóng gói dữ liệu', level=2)
add_paragraph(doc, 'Một khung Modbus RTU ghi một thanh ghi đơn (Function Code 0x06):')
add_formula(doc,
    '| Slave ID | Func Code | Reg Address | Data Value | CRC |'
)
add_paragraph(doc,
    'CRC đảm bảo biến tần phát hiện khung bị lỗi do nhiễu và từ chối xử lý. '
    'Thư viện ModbusMaster xử lý tự động việc đóng gói và kiểm tra CRC.'
)

add_heading(doc, '7.4 Retry logic trong VFDManager', level=2)
add_paragraph(doc,
    'Mỗi lệnh ghi được thử tối đa 3 lần. Nếu cả 3 lần thất bại mới báo lỗi. '
    'Điều này tăng độ tin cậy trong môi trường có nhiễu điện từ từ biến tần và các thiết bị lân cận.'
)

# ─────────────────────────────────────────────
# CHƯƠNG 8 – DATA STORAGE
# ─────────────────────────────────────────────

add_heading(doc, '8. DATA STORAGE (Lưu cấu hình)')

add_paragraph(doc,
    'Toàn bộ tham số vận hành được lưu vào vùng nhớ Flash theo cơ chế NVS '
    '(Non-Volatile Storage) của ESP32. Dữ liệu tồn tại sau khi tắt nguồn, '
    'không cần EEPROM rời, đọc/ghi từ firmware dễ dàng.'
)

add_heading(doc, '8.1 Cấu trúc tham số lưu trữ', level=2)
add_table(doc,
    headers=['Key NVS', 'Tham số', 'Mô tả'],
    rows=[
        ['uen', 'Cho phép tự động', 'true/false'],
        ['udmin / udmax', 'Khoảng cách min/max', 'cm – biên mapping'],
        ['ufmin / ufmax', 'Tần số min/max', 'Hz – biên tần số'],
        ['b1_pct / b2_pct / b3_pct', 'Phần trăm BOOST cấp 1/2/3', '%'],
        ['b1_hold / b2_hold / b3_hold', 'Thời gian giữ BOOST', 'giây'],
        ['b_esc2 / b_esc3', 'Thời gian leo cấp 2/3', 'giây'],
        ['b_decay', 'Thời gian suy giảm mỗi cấp', 'giây'],
        ['wifi_ssid / wifi_pass', 'Cấu hình Wi-Fi', 'String'],
    ],
    col_widths=[4.5, 5.0, 8.0]
)

add_heading(doc, '8.2 Quy trình load khi khởi động', level=2)
add_paragraph(doc, 'Khi firmware khởi động, quá trình đọc cấu hình thực hiện theo thứ tự:')
add_bullet(doc, 'Mở vùng nhớ NVS namespace "MS".')
add_bullet(doc, 'Dọn các key cũ đã bị loại bỏ khỏi firmware hiện tại.')
add_bullet(doc, 'Kiểm tra từng tham số — nếu thiếu thì dùng giá trị mặc định an toàn.')
add_bullet(doc, 'Ràng buộc lại biên (sanitize) để tránh cấu hình vô lý.')
add_bullet(doc, 'Nếu có tham số mới chưa được ghi, tự động lưu lại giá trị mặc định.')
add_note(doc,
    'Nếu người dùng cấu hình tần số lớn nhất nhỏ hơn tần số nhỏ nhất, '
    'firmware tự hiệu chỉnh để hệ thống vẫn hoạt động an toàn.'
)

# ─────────────────────────────────────────────
# CHƯƠNG 9 – USER INTERFACE
# ─────────────────────────────────────────────

add_heading(doc, '9. USER INTERFACE (Giao diện người dùng)')

add_heading(doc, '9.1 Màn hình OLED SH1106', level=2)
add_paragraph(doc,
    'Màn hình OLED 128×64 pixel đóng vai trò HMI tại chỗ, '
    'hiển thị liên tục các thông số vận hành. Nội dung màn hình:'
)
add_table(doc,
    headers=['Vùng hiển thị', 'Nội dung', 'Ghi chú'],
    rows=[
        ['Header', 'MODE: AUTO / TEST / STOP', 'Trạng thái chế độ hiện tại'],
        ['Trung tâm trái', 'Distance: xx.x cm', 'Khoảng cách đo được'],
        ['Trung tâm phải', 'Boosting: ON / OFF', 'Trạng thái BOOST'],
        ['Dưới giữa', 'Target F: xx.x Hz', 'Tần số mục tiêu'],
        ['Footer trái', 'B: ON/OFF', 'Cảm biến thiếu dây'],
        ['Footer phải', 'VFD: xx.xHz', 'Tần số đã gửi xuống VFD'],
        ['Icon nhấp nháy', '>>> khi gửi', 'Chỉ hiện khi task Modbus đang chạy'],
    ],
    col_widths=[4.0, 5.5, 8.0]
)

add_heading(doc, '9.2 Web HMI nội bộ', level=2)
add_paragraph(doc,
    'Hệ thống cung cấp thêm giao diện web truy cập qua IP nội bộ. '
    'Giao diện HMI công nghiệp: nền đen, màu xanh cyan, font monospace. '
    'Người dùng có thể:'
)
add_bullet(doc, 'Theo dõi real-time: distance, frequency, boost status, boost countdown.')
add_bullet(doc, 'Chỉnh tham số: d_min, d_max, f_min, f_max, boost levels, hold times.')
add_bullet(doc, 'Xoay màn hình OLED 180° ngay từ web.')
add_bullet(doc, 'Dữ liệu được cập nhật mỗi 700 ms bằng fetch API.')

# ─────────────────────────────────────────────
# CHƯƠNG 10 – SYSTEM MODES
# ─────────────────────────────────────────────

add_heading(doc, '10. SYSTEM MODES (Chế độ vận hành)')
add_table(doc,
    headers=['Chế độ', 'Lệnh kích hoạt', 'Mô tả', 'Ứng dụng'],
    rows=[
        ['AUTO (MODE_RUN)',
         '"RUN" qua Serial',
         'Hệ thống tự đọc cảm biến, tính tần số, gửi lệnh VFD.',
         'Vận hành sản xuất liên tục.'],
        ['MANUAL (MODE_MANUAL)',
         '"STOP" qua Serial',
         'Dừng VFD, cho phép can thiệp thủ công.',
         'Bảo trì, điều chỉnh cơ khí.'],
        ['TEST/PING (MODE_PING)',
         '"PING" qua Serial',
         'Hiển thị dữ liệu cảm biến, không điều khiển VFD.',
         'Commissioning, kiểm tra cảm biến.'],
    ],
    col_widths=[3.5, 4.0, 6.5, 4.5]
)
add_note(doc,
    'TEST là chế độ "nghe ngóng và quan sát", còn AUTO là chế độ "tự lái".'
)

# ─────────────────────────────────────────────
# CHƯƠNG 11 – PERFORMANCE ANALYSIS
# ─────────────────────────────────────────────

add_heading(doc, '11. PERFORMANCE ANALYSIS (Phân tích hiệu suất)')

add_heading(doc, '11.1 So sánh trước và sau cải tiến', level=2)
add_table(doc,
    headers=['Tiêu chí', 'Điều khiển thủ công', 'Hệ thống tự động'],
    rows=[
        ['Phụ thuộc operator', 'Cao', 'Thấp rõ rệt'],
        ['Tính đồng đều giữa các ca', 'Không ổn định', 'Ổn định theo cấu hình'],
        ['Phản ứng khi thiếu dây', 'Phụ thuộc người quan sát', 'Tự kích BOOST gần như tức thời'],
        ['Lặp lại tham số', 'Kém', 'Tốt – lưu Flash'],
        ['Giám sát tại chỗ', 'Hạn chế', 'OLED + Web HMI'],
        ['Khả năng mở rộng', 'Thấp', 'Cao – modular code'],
    ],
    col_widths=[5.5, 5.5, 6.5]
)

add_heading(doc, '11.2 Độ ổn định', level=2)
add_paragraph(doc,
    'Độ ổn định đến từ nhiều lớp kỹ thuật kết hợp: '
    'lọc EMA cho khoảng cách, debounce cảm biến số (40 ms), '
    'vùng chết 0.2 Hz chống spam lệnh, cơ chế decay BOOST và '
    'task nền Modbus độc lập. Nhờ vậy motor không bị đổi tốc độ dồn dập, '
    'hệ thống ít dao động và "mượt" hơn rõ rệt so với chỉnh tay liên tục.'
)

add_heading(doc, '11.3 Độ phản hồi', level=2)
add_paragraph(doc,
    'Vòng lặp điều khiển chạy mỗi 20 ms (~50 Hz). '
    'Lệnh Modbus được cập nhật mỗi 200 ms (UPDATE_CYCLE). '
    'Đây là sự cân bằng tốt giữa phản ứng đủ nhanh với thay đổi dây '
    'và tránh overload bus truyền thông. '
    'Đối với hệ cơ khí có quán tính như drum và motor kéo dây, '
    'tần số cập nhật 5 Hz là hoàn toàn phù hợp.'
)

# ─────────────────────────────────────────────
# CHƯƠNG 12 – LIMITATIONS
# ─────────────────────────────────────────────

add_heading(doc, '12. LIMITATIONS (Giới hạn)')

add_bullet(doc,
    'Nhiễu cảm biến siêu âm: phụ thuộc bề mặt phản xạ, góc lắp và môi trường rung động mạnh.'
)
add_bullet(doc,
    'Độ trễ Modbus: RS485 không siêu nhanh; nếu bus nhiễu hoặc VFD phản hồi chậm, lệnh có độ trễ nhất định.'
)
add_bullet(doc,
    'Khoảng cách là đại lượng gián tiếp: chưa phải lực căng thực của dây, '
    'do đó thuật toán điều khiển theo "trạng thái hình học" chứ chưa phải closed-loop lực căng.'
)
add_bullet(doc,
    'MANUAL chưa phải closed-loop độc lập: thiên về phục vụ bảo trì, '
    'chưa có nhánh điều khiển tốc độ thủ công đầy đủ (HMI slider).'
)
add_bullet(doc,
    'Một cảm biến duy nhất: nếu cảm biến hỏng sẽ mất toàn bộ dữ liệu đầu vào.'
)

# ─────────────────────────────────────────────
# CHƯƠNG 13 – FUTURE IMPROVEMENTS
# ─────────────────────────────────────────────

add_heading(doc, '13. FUTURE IMPROVEMENTS (Cải tiến tương lai)')

add_table(doc,
    headers=['Hướng cải tiến', 'Mô tả', 'Lợi ích'],
    rows=[
        ['PID Controller',
         'Thay mapping tuyến tính bằng PID để bám theo giá trị setpoint.',
         'Giảm sai số, phản ứng tốt hơn khi tải thay đổi.'],
        ['Kalman Filter',
         'Lọc khoảng cách chính xác hơn trong môi trường nhiễu cao.',
         'Ước lượng trạng thái mượt hơn EMA đơn giản.'],
        ['Multi-sensor Fusion',
         'Kết hợp siêu âm + dancer sensor + load cell + encoder.',
         'Tăng độ tin cậy, không phụ thuộc một cảm biến.'],
        ['IoT Monitoring',
         'Dashboard từ xa, log sản xuất, cảnh báo lỗi.',
         'Giám sát nhiều máy, phân tích xu hướng.'],
        ['Fail-safe Logic',
         'Watchdog, cảnh báo mất tín hiệu, mode an toàn khi lỗi cảm biến.',
         'Tăng độ tin cậy trong môi trường công nghiệp thực tế.'],
        ['Recipe Management',
         'Nhiều bộ tham số cho từng sản phẩm, chuyển đổi nhanh.',
         'Linh hoạt cho dây chuyền sản xuất nhiều chủng loại.'],
    ],
    col_widths=[4.0, 7.0, 6.5]
)

# ─────────────────────────────────────────────
# CHƯƠNG 14 – CONCLUSION
# ─────────────────────────────────────────────

add_heading(doc, '14. CONCLUSION (Kết luận)')

add_paragraph(doc,
    'Hệ thống điều khiển tốc độ motor qua biến tần sử dụng ESP32-S3, cảm biến siêu âm '
    'và Modbus RTU là một giải pháp thực tiễn, có giá trị triển khai cao trong môi trường '
    'sản xuất công nghiệp. Điểm mạnh không chỉ nằm ở việc tự động thay đổi tần số theo '
    'trạng thái drum, mà còn ở khả năng tổ chức bài bản: BOOST bù tốc nhiều cấp, '
    'lưu cấu hình bền vững, hiển thị trực quan và nền tảng mở rộng.'
)
add_paragraph(doc,
    'Về mặt kỹ thuật, đề tài đã giải quyết được bài toán cốt lõi là biến thông tin '
    'cơ khí thành lệnh điều khiển điện một cách logic và có kiểm soát. '
    'Về mặt vận hành, hệ thống giúp giảm phụ thuộc vào kinh nghiệm cá nhân, '
    'tăng tính ổn định quá trình cấp dây và tạo tiền đề cho chuyển đổi sang sản xuất thông minh hơn.'
)
add_paragraph(doc,
    'Tóm lại, đây là một mô hình điều khiển nhúng mang tính ứng dụng thực tế cao, '
    'phù hợp cho các dây chuyền vừa và nhỏ, đồng thời có đủ chiều sâu để phát triển '
    'tiếp lên các cấp độ điều khiển tiên tiến hơn trong tương lai.'
)

add_separator(doc)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run(
    f'ESP32-S3 VFD Controller  |  Modbus RTU + RS485  |  '
    f'Báo cáo kỹ thuật – {datetime.date.today().strftime("%d/%m/%Y")}'
)
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.italic = True

# ─────────────────────────────────────────────
# LƯU FILE
# ─────────────────────────────────────────────

output_path = 'BaoCaoKyThuat_VFD_ESP32.docx'
doc.save(output_path)
print(f'[OK] Da tao file: {output_path}')
